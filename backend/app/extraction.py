"""
Приведение любого вложения к тексту.

Порядок для PDF: сначала текстовый слой (pdfminer.six — быстро и точно),
и только если его нет или он пустой — растеризация и OCR. Это важно: OCR по
PDF с готовым текстовым слоем даёт худший результат и тратит время.

Все внешние инструменты опциональны. Если tesseract или poppler не установлены,
модуль не падает: вложение сохраняется, помечается needs_review, а причина
показывается в интерфейсе, чтобы текст можно было вставить руками.
"""

from __future__ import annotations

import hashlib
import io
import re
import shutil
import subprocess
from dataclasses import dataclass, field
from email import message_from_bytes
from email.header import decode_header, make_header
from pathlib import Path

from .config import settings

TEXT_EXT = {".txt", ".text", ".md", ".csv", ".log", ".json"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif", ".heic"}
HTML_EXT = {".html", ".htm"}
#: Заведомо двоичные форматы — их нельзя «прочитать как текст».
BINARY_EXT = {".zip", ".rar", ".7z", ".gz", ".tar", ".exe", ".dll", ".bin", ".iso",
              ".mp3", ".mp4", ".avi", ".mov", ".wav", ".ppt", ".pptx", ".odp"}
#: Старые бинарные форматы Microsoft Office — нужен пересохранённый файл.
LEGACY_OFFICE_EXT = {".doc": "DOCX", ".xls": "XLSX", ".rtf": "DOCX или PDF"}

#: Ниже этого числа символов на страницу считаем, что текстового слоя нет.
MIN_CHARS_PER_PAGE = 40


@dataclass
class ExtractionResult:
    text: str = ""
    method: str = ""
    error: str = ""
    page_count: int = 0
    needs_review: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text or "")


# --------------------------------------------------------------------------- #
#  Доступность внешних инструментов
# --------------------------------------------------------------------------- #

def tool_status() -> dict:
    """Показывается в интерфейсе, чтобы отсутствие OCR не было сюрпризом."""
    tesseract = shutil.which("tesseract")
    poppler = shutil.which("pdftoppm")
    langs: list[str] = []
    if tesseract:
        try:
            out = subprocess.run([tesseract, "--list-langs"], capture_output=True,
                                 text=True, timeout=15)
            langs = [ln.strip() for ln in out.stdout.splitlines()[1:] if ln.strip()]
        except Exception:
            pass
    return {
        "tesseract": bool(tesseract),
        "tesseract_path": tesseract or "",
        "tesseract_langs": langs,
        "russian_ocr": "rus" in langs,
        "poppler": bool(poppler),
        "pdf_text": True,
        "docx": True,
        "configured_lang": settings.tesseract_lang,
    }


# --------------------------------------------------------------------------- #
#  Помощники
# --------------------------------------------------------------------------- #

_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_MULTI_NL = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def clean_text(text: str) -> str:
    t = _CTRL.sub("", text or "")
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # Склейка переносов по дефису, типичных для OCR и PDF.
    t = re.sub(r"(\w)-\n(\w)", r"\1\2", t)
    t = _TRAILING_WS.sub("\n", t)
    t = _MULTI_NL.sub("\n\n", t)
    return t.strip()


def decode_bytes(data: bytes) -> str:
    """Определить кодировку: важно для писем в windows-1251 и koi8-r."""
    for enc in ("utf-8", "utf-8-sig"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            pass
    try:
        import chardet
        guess = chardet.detect(data)
        if guess.get("encoding") and (guess.get("confidence") or 0) > 0.5:
            return data.decode(guess["encoding"], errors="replace")
    except Exception:
        pass
    for enc in ("cp1251", "koi8-r", "iso-8859-5"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def strip_html(html: str) -> str:
    from html.parser import HTMLParser

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts: list[str] = []
            self.skip = 0

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style"):
                self.skip += 1
            elif tag in ("br", "p", "div", "tr", "li", "h1", "h2", "h3", "table"):
                self.parts.append("\n")

        def handle_endtag(self, tag):
            if tag in ("script", "style") and self.skip:
                self.skip -= 1

        def handle_data(self, data):
            if not self.skip:
                self.parts.append(data)

    p = _P()
    try:
        p.feed(html)
    except Exception:
        return clean_text(re.sub(r"<[^>]+>", " ", html))
    return clean_text("".join(p.parts))


def sha256_of(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# --------------------------------------------------------------------------- #
#  Извлечение по типам
# --------------------------------------------------------------------------- #

def extract_pdf(data: bytes) -> ExtractionResult:
    res = ExtractionResult()
    pages = 0
    text = ""

    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(data)) or ""
    except Exception as exc:
        res.warnings.append(f"Текстовый слой не прочитан: {exc}")

    try:
        from pypdf import PdfReader
        pages = len(PdfReader(io.BytesIO(data)).pages)
    except Exception:
        pages = 0
    res.page_count = pages

    cleaned = clean_text(text)
    threshold = MIN_CHARS_PER_PAGE * max(pages, 1)
    if len(cleaned) >= threshold:
        res.text = cleaned
        res.method = "PDF_TEXT"
        return res

    # Текстового слоя нет — это скан. Пробуем OCR.
    ocr = ocr_pdf(data)
    if ocr.text:
        if cleaned:
            ocr.warnings.append(
                f"Текстовый слой содержал всего {len(cleaned)} символов на {pages} стр. — "
                f"документ распознан как скан."
            )
        ocr.page_count = pages or ocr.page_count
        return ocr

    res.text = cleaned
    res.method = "PDF_TEXT" if cleaned else "FAILED"
    res.needs_review = True
    res.error = ocr.error or "PDF не содержит текстового слоя, OCR недоступен."
    res.warnings.extend(ocr.warnings)
    return res


def ocr_pdf(data: bytes) -> ExtractionResult:
    res = ExtractionResult(method="PDF_OCR")
    if not shutil.which("tesseract"):
        res.error = ("Не установлен tesseract-ocr — распознавание сканов недоступно. "
                     "Установите: apt-get install tesseract-ocr tesseract-ocr-rus")
        res.needs_review = True
        return res
    if not shutil.which("pdftoppm"):
        res.error = ("Не установлен poppler-utils — PDF нельзя растеризовать для OCR. "
                     "Установите: apt-get install poppler-utils")
        res.needs_review = True
        return res
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(data, dpi=settings.ocr_dpi)
        res.page_count = len(images)
        chunks: list[str] = []
        for i, img in enumerate(images, 1):
            page = pytesseract.image_to_string(img, lang=settings.tesseract_lang)
            chunks.append(f"--- страница {i} ---\n{page.strip()}")
        res.text = clean_text("\n\n".join(chunks))
        res.needs_review = len(res.text) < 80
        if res.needs_review:
            res.warnings.append("OCR вернул очень мало текста — проверьте качество скана.")
    except Exception as exc:
        res.error = f"Ошибка OCR: {exc}"
        res.needs_review = True
    return res


def extract_image(data: bytes) -> ExtractionResult:
    res = ExtractionResult(method="OCR", page_count=1)
    if not shutil.which("tesseract"):
        res.error = ("Не установлен tesseract-ocr — распознавание фотографий недоступно. "
                     "Установите: apt-get install tesseract-ocr tesseract-ocr-rus")
        res.needs_review = True
        return res
    try:
        import pytesseract
        from PIL import Image, ImageOps

        img = Image.open(io.BytesIO(data))
        img = ImageOps.exif_transpose(img)          # учесть поворот из EXIF
        if img.mode not in ("L", "RGB"):
            img = img.convert("RGB")
        # Апскейл мелких снимков заметно улучшает распознавание.
        if min(img.size) < 1000:
            k = 1000 / max(min(img.size), 1)
            img = img.resize((int(img.width * k), int(img.height * k)), Image.LANCZOS)
        img = ImageOps.grayscale(img)
        img = ImageOps.autocontrast(img)
        res.text = clean_text(pytesseract.image_to_string(img, lang=settings.tesseract_lang))
        res.needs_review = len(res.text) < 40
        if res.needs_review:
            res.warnings.append(
                "Распознано мало текста. Проверьте, что фото чёткое и не перевёрнуто, "
                "либо вставьте текст вручную."
            )
    except Exception as exc:
        res.error = f"Ошибка OCR изображения: {exc}"
        res.needs_review = True
    return res


def extract_docx(data: bytes) -> ExtractionResult:
    res = ExtractionResult(method="DOCX")
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        res.text = clean_text("\n".join(parts))
    except Exception as exc:
        res.error = f"Не удалось прочитать .docx: {exc}"
        res.needs_review = True
    return res


def extract_xlsx(data: bytes) -> ExtractionResult:
    res = ExtractionResult(method="XLSX")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"--- лист: {ws.title} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        res.text = clean_text("\n".join(parts))
    except Exception as exc:
        res.error = f"Не удалось прочитать .xlsx: {exc}"
        res.needs_review = True
    return res


def extract_eml(data: bytes) -> ExtractionResult:
    """Разбор .eml/.msg-подобного письма вместе с вложениями."""
    res = ExtractionResult(method="EML")
    try:
        msg = message_from_bytes(data)

        def hdr(name: str) -> str:
            raw = msg.get(name)
            if not raw:
                return ""
            try:
                return str(make_header(decode_header(raw)))
            except Exception:
                return raw

        head = [
            f"От: {hdr('From')}", f"Кому: {hdr('To')}", f"Дата: {hdr('Date')}",
            f"Тема: {hdr('Subject')}",
        ]
        body_parts: list[str] = []
        html_parts: list[str] = []
        nested: list[str] = []

        for part in msg.walk():
            if part.get_content_maintype() == "multipart":
                continue
            disp = (part.get("Content-Disposition") or "").lower()
            ctype = part.get_content_type()
            payload = part.get_payload(decode=True) or b""
            if "attachment" in disp:
                fname = part.get_filename() or "вложение"
                try:
                    fname = str(make_header(decode_header(fname)))
                except Exception:
                    pass
                sub = extract(fname, payload)
                nested.append(f"--- вложение письма: {fname} ({sub.method}) ---\n{sub.text}")
                continue
            if ctype == "text/plain":
                body_parts.append(decode_bytes(payload))
            elif ctype == "text/html":
                html_parts.append(strip_html(decode_bytes(payload)))

        body = "\n".join(body_parts) or "\n".join(html_parts)
        res.text = clean_text("\n".join(head) + "\n\n" + body + "\n\n" + "\n\n".join(nested))
    except Exception as exc:
        res.error = f"Не удалось разобрать письмо: {exc}"
        res.needs_review = True
    return res


def extract(filename: str, data: bytes) -> ExtractionResult:
    """Диспетчер по расширению и сигнатуре файла."""
    ext = Path(filename or "").suffix.lower()

    if data[:5] == b"%PDF-" or ext == ".pdf":
        return extract_pdf(data)
    if ext == ".docx" or (data[:2] == b"PK" and ext in (".docx", ".dotx")):
        return extract_docx(data)
    if ext in (".xlsx", ".xlsm"):
        return extract_xlsx(data)
    if ext in (".eml", ".msg"):
        return extract_eml(data)
    if ext in IMAGE_EXT or data[:4] in (b"\x89PNG", b"\xff\xd8\xff\xe0", b"\xff\xd8\xff\xe1") \
            or data[:3] == b"\xff\xd8\xff":
        return extract_image(data)
    if ext in HTML_EXT:
        return ExtractionResult(text=strip_html(decode_bytes(data)), method="HTML")
    if ext in TEXT_EXT or not ext:
        text = clean_text(decode_bytes(data))
        return ExtractionResult(text=text, method="PLAIN", needs_review=not text)

    if ext in LEGACY_OFFICE_EXT:
        return ExtractionResult(
            method="FAILED", needs_review=True,
            error=f"Формат «{ext}» устарел и не читается напрямую. Пересохраните файл "
                  f"как {LEGACY_OFFICE_EXT[ext]} либо вставьте текст вручную.",
        )
    if ext in BINARY_EXT or data[:2] == b"PK" or data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return ExtractionResult(
            method="FAILED", needs_review=True,
            error=f"Файл «{ext or 'без расширения'}» — двоичный, текст из него не "
                  f"извлекается. Приложите документ в PDF, DOCX или изображении.",
        )

    # Неизвестный тип: читаем как текст, только если это действительно похоже на текст.
    text = clean_text(decode_bytes(data))
    sample = text[:4000]
    if len(sample) >= 40:
        printable = sum(ch.isprintable() or ch in "\n\t" for ch in sample)
        letters = sum(ch.isalpha() or ch.isspace() for ch in sample)
        if printable / len(sample) > 0.95 and letters / len(sample) > 0.6:
            return ExtractionResult(
                text=text, method="PLAIN",
                warnings=[f"Неизвестный тип «{ext}», прочитан как обычный текст."])
    return ExtractionResult(
        method="FAILED", needs_review=True,
        error=f"Формат «{ext or 'без расширения'}» не поддерживается. "
              f"Сохраните документ в PDF, DOCX или вставьте текст вручную.",
    )


def save_upload(filename: str, data: bytes) -> tuple[Path, str]:
    """Сохранить файл в хранилище, дедуплицируя по SHA-256."""
    digest = sha256_of(data)
    root = Path(settings.upload_dir) / digest[:2]
    root.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^\w.\-]+", "_", Path(filename or "file").name)[:120] or "file"
    path = root / f"{digest[:16]}_{safe}"
    if not path.exists():
        path.write_bytes(data)
    return path, digest
